import os
import logging
from PIL import Image
from pdf2image import convert_from_path
from PyPDF2 import PdfMerger
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import List, Dict, Any
import tempfile
from datetime import datetime

class ImageProcessor:
    def __init__(self):
        self.temp_files = []
    
    def pdf_to_png(self, pdf_path: str, dpi: int = 200) -> str:
        """Convert PDF to PNG image"""
        try:
            # Convert PDF to images
            images = convert_from_path(pdf_path, dpi=dpi)
            
            if not images:
                raise Exception("No images found in PDF")
            
            # Use the first page
            image = images[0]
            
            # Generate PNG filename
            base_name = os.path.splitext(os.path.basename(pdf_path))[0]
            png_path = os.path.join(os.path.dirname(pdf_path), f"{base_name}.png")
            
            # Save as PNG
            image.save(png_path, "PNG")
            
            logging.info(f"Successfully converted PDF to PNG: {png_path}")
            return png_path
            
        except Exception as e:
            logging.error(f"Failed to convert PDF to PNG: {str(e)}")
            raise Exception(f"PDF conversion failed: {str(e)}")
    
    def crop_image(self, image_path: str, crop_data: Dict[str, float]) -> str:
        """Crop an image based on crop coordinates"""
        try:
            image = Image.open(image_path)
            
            # Extract crop coordinates
            x1 = int(crop_data['x'])
            y1 = int(crop_data['y'])
            x2 = int(crop_data['x'] + crop_data['width'])
            y2 = int(crop_data['y'] + crop_data['height'])
            
            # Ensure coordinates are within image bounds
            x1 = max(0, min(x1, image.width))
            y1 = max(0, min(y1, image.height))
            x2 = max(0, min(x2, image.width))
            y2 = max(0, min(y2, image.height))
            
            # Ensure we have a valid crop area
            if x2 <= x1 or y2 <= y1:
                raise Exception("Invalid crop coordinates")
            
            # Crop the image
            cropped_image = image.crop((x1, y1, x2, y2))
            
            # Generate cropped filename
            base_name = os.path.splitext(os.path.basename(image_path))[0]
            cropped_path = os.path.join(os.path.dirname(image_path), f"{base_name}_cropped.png")
            
            # Save cropped image
            cropped_image.save(cropped_path, "PNG")
            
            logging.info(f"Successfully cropped image: {cropped_path}")
            return cropped_path
            
        except Exception as e:
            logging.error(f"Failed to crop image: {str(e)}")
            raise Exception(f"Image cropping failed: {str(e)}")
    
    def combine_to_pdf(self, image_paths: List[str], output_dir: str, filename: str) -> str:
        """Combine multiple images into a single PDF"""
        try:
            output_path = os.path.join(output_dir, f"{filename}.pdf")
            
            # Convert images to PDF
            temp_pdfs = []
            merger = PdfMerger()
            
            for i, image_path in enumerate(image_paths):
                if not os.path.exists(image_path):
                    logging.warning(f"Image not found: {image_path}")
                    continue
                
                # Open and convert image to RGB if necessary
                image = Image.open(image_path)
                if image.mode != 'RGB':
                    image = image.convert('RGB')
                
                # Create temporary PDF for this image
                temp_pdf_path = os.path.join(output_dir, f"temp_{i}.pdf")
                image.save(temp_pdf_path, "PDF")
                temp_pdfs.append(temp_pdf_path)
                
                # Add to merger
                merger.append(temp_pdf_path)
            
            if not temp_pdfs:
                raise Exception("No valid images to combine")
            
            # Write combined PDF
            merger.write(output_path)
            merger.close()
            
            # Clean up temporary PDFs
            for temp_pdf in temp_pdfs:
                try:
                    os.remove(temp_pdf)
                except:
                    pass
            
            logging.info(f"Successfully created combined PDF: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"Failed to combine images to PDF: {str(e)}")
            raise Exception(f"PDF combination failed: {str(e)}")
    
    def combine_to_word(self, image_paths: List[str], output_dir: str, filename: str) -> str:
        """Combine multiple images into a single Word document"""
        try:
            output_path = os.path.join(output_dir, f"{filename}.docx")
            
            # Create new Word document
            doc = Document()
            doc.add_heading('Tableau Dashboard Export', 0)
            
            for i, image_path in enumerate(image_paths):
                if not os.path.exists(image_path):
                    logging.warning(f"Image not found: {image_path}")
                    continue
                
                # Add section heading
                doc.add_heading(f'Dashboard {i + 1}', level=1)
                
                # Add image to document
                # Calculate appropriate width (max 6 inches)
                image = Image.open(image_path)
                aspect_ratio = image.height / image.width
                width = min(6.0, image.width / 100)  # Convert pixels to inches roughly
                height = width * aspect_ratio
                
                doc.add_picture(image_path, width=Inches(width))
                
                # Add page break if not the last image
                if i < len(image_paths) - 1:
                    doc.add_page_break()
            
            # Save document
            doc.save(output_path)
            
            logging.info(f"Successfully created Word document: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"Failed to combine images to Word: {str(e)}")
            raise Exception(f"Word document creation failed: {str(e)}")
    
    def combine_to_word_with_details(self, image_paths: List[str], output_dir: str, filename: str, summary_data: List[Dict]) -> str:
        """Combine multiple images into a single Word document with detailed metadata using 2-column layout"""
        try:
            output_path = os.path.join(output_dir, f"{filename}.docx")
            
            # Create new Word document
            doc = Document()
            
            # Add main title
            title = doc.add_heading('Tableau Dashboard Export Report', 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add summary information on first page
            doc.add_heading('Export Summary', level=1)
            summary_para = doc.add_paragraph()
            summary_para.add_run('Generated: ').bold = True
            summary_para.add_run(f'{datetime.now().strftime("%Y-%m-%d %H:%M:%S")}\n')
            summary_para.add_run('Total Dashboards: ').bold = True
            summary_para.add_run(f'{len(summary_data)}\n')
            summary_para.add_run('Export Format: ').bold = True
            summary_para.add_run('Microsoft Word Document (.docx)')
            
            # Add page break after summary
            doc.add_page_break()
            
            # Process dashboards in pairs for same-page layout
            for i in range(0, len(image_paths), 2):
                # Add page heading for dashboard(s)
                if i + 1 < len(image_paths):
                    page_title = f'Dashboards {i + 1} & {i + 2}'
                else:
                    page_title = f'Dashboard {i + 1}'
                
                doc.add_heading(page_title, level=1)
                
                # First dashboard
                self._add_dashboard_to_word(doc, image_paths[i], summary_data[i], i + 1)
                
                # Second dashboard on same page (if exists)
                if i + 1 < len(image_paths):
                    # Add some spacing between dashboards
                    doc.add_paragraph()
                    self._add_dashboard_to_word(doc, image_paths[i + 1], summary_data[i + 1], i + 2)
                
                # Add page break if not the last pair
                if i + 2 < len(image_paths):
                    doc.add_page_break()
            
            # Save document
            doc.save(output_path)
            
            logging.info(f"Successfully created detailed Word document: {output_path}")
            return output_path
            
        except Exception as e:
            logging.error(f"Failed to combine images to Word with details: {str(e)}")
            raise Exception(f"Detailed Word document creation failed: {str(e)}")
    
    def _add_dashboard_to_word(self, doc, image_path: str, data: Dict, section_num: int):
        """Add a single dashboard to Word document with 2-column layout"""
        try:
            if not os.path.exists(image_path):
                logging.warning(f"Image not found: {image_path}")
                return
            
            # Create a table for 2-column layout
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # Set column widths (50% each)
            for cell in table.rows[0].cells:
                cell.width = Inches(3.25)  # Half of 6.5 inch page width
            
            # Left column - Image
            left_cell = table.rows[0].cells[0]
            left_para = left_cell.paragraphs[0]
            
            # Add image to left cell
            image = Image.open(image_path)
            img_width = image.width
            img_height = image.height
            aspect_ratio = img_height / img_width
            
            # Set image width to fit in left column (3 inches max)
            img_width_inches = 3.0
            
            run = left_para.runs[0] if left_para.runs else left_para.add_run()
            run.add_picture(image_path, width=Inches(img_width_inches))
            
            # Right column - Metadata
            right_cell = table.rows[0].cells[1]
            right_para = right_cell.paragraphs[0]
            
            # Add section title
            title_run = right_para.add_run(f'Dashboard {section_num}\n')
            title_run.font.size = Pt(14)
            title_run.bold = True
            
            # Add metadata
            right_para.add_run('\nProject: ').bold = True
            right_para.add_run(f'{data.get("project", "Unknown")}\n')
            
            right_para.add_run('Workbook: ').bold = True
            right_para.add_run(f'{data.get("workbook", "Unknown")}\n')
            
            right_para.add_run('Dashboard: ').bold = True
            right_para.add_run(f'{data.get("dashboard", "Unknown")}\n')
            
            right_para.add_run('Exported: ').bold = True
            right_para.add_run(f'{data.get("timestamp", "Unknown")}')
            
            # Set vertical alignment for cells
            left_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            right_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
            logging.info(f"Added dashboard {section_num} to Word document in 2-column layout")
            
        except Exception as img_error:
            logging.error(f"Failed to add dashboard {section_num}: {str(img_error)}")
            # Add error message instead
            error_para = doc.add_paragraph()
            error_para.add_run(f'[Error loading Dashboard {section_num}: {os.path.basename(image_path)}]').italic = True
    
    def create_thumbnail(self, image_path: str, max_width: int = 200, max_height: int = 120) -> str:
        """Create a thumbnail of an image"""
        try:
            image = Image.open(image_path)
            
            # Calculate thumbnail size maintaining aspect ratio
            image.thumbnail((max_width, max_height), Image.Resampling.LANCZOS)
            
            # Generate thumbnail filename
            base_name = os.path.splitext(os.path.basename(image_path))[0]
            thumb_path = os.path.join(os.path.dirname(image_path), f"{base_name}_thumb.png")
            
            # Save thumbnail
            image.save(thumb_path, "PNG")
            
            logging.info(f"Successfully created thumbnail: {thumb_path}")
            return thumb_path
            
        except Exception as e:
            logging.error(f"Failed to create thumbnail: {str(e)}")
            raise Exception(f"Thumbnail creation failed: {str(e)}")
    
    def cleanup_temp_files(self):
        """Clean up any temporary files created during processing"""
        for temp_file in self.temp_files:
            try:
                if os.path.exists(temp_file):
                    os.remove(temp_file)
            except:
                pass
        self.temp_files.clear()
